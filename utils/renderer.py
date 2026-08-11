import html
from io import BytesIO
import re
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from models.schemas import RenderConfig
from utils.logger import logger
from utils.parser import is_header_line

# 導入 ReportLab 核心套件
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate

# 🎯 註冊 ReportLab 內建 CJK 中文字型 (免安裝外部字型，全平台/雲端通用)
try:
  pdfmetrics.registerFont(UnicodeCIDFont("STHeiti-Light"))
  PDF_FONT = "STHeiti-Light"
except Exception as e:
  logger.warning("無法註冊 CJK 字型，降級為 Helvetica: %s", str(e))
  PDF_FONT = "Helvetica"


def create_docx(raw_text: str, config: RenderConfig) -> bytes:
  logger.info("正生成 Word (.docx) 檔案...")
  doc = docx.Document()

  for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

  style = doc.styles["Normal"]
  font = style.font
  font.name = config.font_name
  font.size = Pt(config.font_size)
  font.color.rgb = RGBColor(0x33, 0x33, 0x33)

  for line in raw_text.splitlines():
    line_str = line.strip()
    if not line_str:
      continue

    if is_header_line(line_str):
      p = doc.add_paragraph()
      p.paragraph_format.space_before = Pt(14)
      p.paragraph_format.space_after = Pt(4)
      p.paragraph_format.keep_with_next = True

      run = p.add_run(line_str.upper())
      run.bold = True
      run.font.size = Pt(config.font_size + 2.5)
      run.font.color.rgb = RGBColor(*config.primary_color_rgb)

      if line_str.upper() in ["RESUME", "CURRICULUM VITAE", "履歷"]:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(config.font_size + 5.5)
      continue

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

    if (
        line_str.startswith("➢")
        or line_str.startswith("-")
        or line_str.startswith("•")
    ):
      p.paragraph_format.left_indent = Inches(0.25)
      clean_item = re.sub(r"^[➢\-•]\s*", "• ", line_str)
      p.add_run(clean_item)
    elif ":" in line_str and len(line_str.split(":")[0]) < 25:
      parts = line_str.split(":", 1)
      r_label = p.add_run(parts[0] + ": ")
      r_label.bold = True
      r_label.font.color.rgb = RGBColor(*config.primary_color_rgb)
      p.add_run(parts[1].strip())
    elif "：" in line_str and len(line_str.split("：")[0]) < 25:
      parts = line_str.split("：", 1)
      r_label = p.add_run(parts[0] + "：")
      r_label.bold = True
      r_label.font.color.rgb = RGBColor(*config.primary_color_rgb)
      p.add_run(parts[1].strip())
    else:
      p.add_run(line_str)

  buffer = BytesIO()
  doc.save(buffer)
  return buffer.getvalue()


def create_pdf(raw_text: str, config: RenderConfig) -> bytes:
  logger.info("正使用 ReportLab (CJK 支援) 生成 PDF 檔案...")
  buffer = BytesIO()

  doc = SimpleDocTemplate(
      buffer,
      pagesize=A4,
      leftMargin=40,
      rightMargin=40,
      topMargin=40,
      bottomMargin=40,
  )

  story = []
  styles = getSampleStyleSheet()

  hex_color = config.primary_color_hex.lstrip("#")
  r, g, b = tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
  brand_color = colors.Color(r / 255.0, g / 255.0, b / 255.0)

  # 套用中英文兼容字型 PDF_FONT
  title_style = ParagraphStyle(
      "DocTitle",
      parent=styles["Normal"],
      fontName=PDF_FONT,
      fontSize=config.font_size + 6,
      leading=config.font_size + 8,
      textColor=brand_color,
      alignment=1,
      spaceAfter=12,
  )

  header_style = ParagraphStyle(
      "SectionHeader",
      parent=styles["Normal"],
      fontName=PDF_FONT,
      fontSize=config.font_size + 2,
      leading=config.font_size + 4,
      textColor=brand_color,
      spaceBefore=12,
      spaceAfter=4,
      keepWithNext=True,
  )

  body_style = ParagraphStyle(
      "BodyText",
      parent=styles["Normal"],
      fontName=PDF_FONT,
      fontSize=config.font_size,
      leading=config.font_size + 3,
      textColor=colors.HexColor("#333333"),
      spaceAfter=4,
  )

  bullet_style = ParagraphStyle(
      "BulletText",
      parent=body_style,
      leftIndent=15,
      firstLineIndent=-10,
      spaceAfter=3,
  )

  for line in raw_text.splitlines():
    line_str = line.strip()
    if not line_str:
      continue

    safe_line = html.escape(line_str)

    if is_header_line(line_str):
      if line_str.upper() in ["RESUME", "CURRICULUM VITAE", "履歷"]:
        story.append(Paragraph(safe_line.upper(), title_style))
      else:
        story.append(Paragraph(safe_line.upper(), header_style))
        story.append(
            HRFlowable(
                width="100%",
                thickness=1,
                color=brand_color,
                spaceBefore=1,
                spaceAfter=6,
            )
        )
      continue

    if (
        line_str.startswith("➢")
        or line_str.startswith("-")
        or line_str.startswith("•")
    ):
      clean_item = re.sub(r"^[➢\-•]\s*", "&bull; ", safe_line)
      story.append(Paragraph(clean_item, bullet_style))
    elif ":" in line_str and len(line_str.split(":")[0]) < 25:
      parts = safe_line.split(":", 1)
      formatted_p = (
          f"<b><font color='{config.primary_color_hex}'>{parts[0]}:</font></b>"
          f" {parts[1].strip()}"
      )
      story.append(Paragraph(formatted_p, body_style))
    elif "：" in line_str and len(line_str.split("：")[0]) < 25:
      parts = safe_line.split("：", 1)
      formatted_p = (
          f"<b><font color='{config.primary_color_hex}'>{parts[0]}：</font></b>"
          f" {parts[1].strip()}"
      )
      story.append(Paragraph(formatted_p, body_style))
    else:
      story.append(Paragraph(safe_line, body_style))

  doc.build(story)
  pdf_bytes = buffer.getvalue()
  buffer.close()
  return pdf_bytes
